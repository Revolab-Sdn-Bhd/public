import argparse
import glob
import hashlib
import json
import logging
import os
import re
import sys
import tempfile
import uuid
from typing import Dict, List, Literal, Union

import fitz
import pandas as pd
import psycopg2
import psycopg2.extras
from docx import Document
from openai import AzureOpenAI
from pgvector.psycopg2 import register_vector
from unstructured.documents.elements import ElementType
from unstructured.partition.pdf import partition_pdf

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - [%(levelname)s] - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)

from urllib.parse import unquote, urlparse

from google.cloud import storage
from google.oauth2 import service_account

SERVICE_ACCOUNT_FILE = os.getenv("GOOGLE_APPLICATION_CREDENTIALS", "")
QDRANT_UPLOAD_BATCH_SIZE = int(os.getenv("QDRANT_UPLOAD_BATCH_SIZE"))
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "")
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY", "")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT", "")
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION", "")
AZURE_OPENAI_DEPLOYMENT_NAME = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4o")
AZURE_OPENAI_EMBEDDING_API_KEY = os.getenv("AZURE_OPENAI_EMBEDDING_API_KEY", "")
AZURE_OPENAI_EMBEDDING_ENDPOINT = os.getenv("AZURE_OPENAI_EMBEDDING_ENDPOINT", "")
AZURE_OPENAI_EMBEDDING_API_VERSION = os.getenv("AZURE_OPENAI_EMBEDDING_API_VERSION", "")
PRE_LOGIN_COLLECTION = os.getenv("PRE_LOGIN_COLLECTION", "")
POST_LOGIN_COLLECTION = os.getenv("POST_LOGIN_COLLECTION", "")
PG_HOST = os.getenv("PG_HOST", "")
PG_PORT = os.getenv("PG_PORT", "")
PG_USER = os.getenv("PG_USER", "")
PG_PASSWORD = os.getenv("PG_PASSWORD", "")
PG_DATABASE = os.getenv("PG_DATABASE", "")


class Knowledge:
    def __init__(
        self,
        *,
        module: Literal["pre_login", "post_login"] = "pre_login",
    ):
        """
        Initialize the Knowledge object with connection to Postgres.
        Instead of a Qdrant URL and API key, we connect to Postgres using the credentials in configs.
        """
        # Embedding configuration
        self.embedding_endpoint = AZURE_OPENAI_EMBEDDING_ENDPOINT
        self.embedding_model_name = EMBEDDING_MODEL

        self.embedding_api_key = AZURE_OPENAI_EMBEDDING_API_KEY
        self.embedding_api_version = AZURE_OPENAI_EMBEDDING_API_VERSION
        self.azure_openai_embedding_client = AzureOpenAI(
            api_key=self.embedding_api_key,
            api_version=self.embedding_api_version,
            azure_endpoint=self.embedding_endpoint,
        )

        self.azure_openai_endpoint = AZURE_OPENAI_ENDPOINT
        self.azure_openai_api_key = AZURE_OPENAI_API_KEY
        self.azure_openai_api_version = AZURE_OPENAI_API_VERSION
        self.azure_openai_deployment_name = AZURE_OPENAI_DEPLOYMENT_NAME
        self.azure_openai_client = AzureOpenAI(
            api_key=self.azure_openai_api_key,
            azure_endpoint=self.azure_openai_endpoint,
            api_version=self.azure_openai_api_version,
        )

        if module not in ["pre_login", "post_login"]:
            raise ValueError("module must be either 'pre_login' or 'post_login'")
        self.module = module

        if self.module == "pre_login":
            self.faq_collection_name = PRE_LOGIN_COLLECTION  # e.g. "pre_login_faqs"
            self.overall_collec = "pre_v3"  # a fallback collection/table name if needed
        elif self.module == "post_login":
            self.faq_collection_name = POST_LOGIN_COLLECTION  # e.g. "post_login_faqs"
            self.overall_collec = "post_v3"

        # Postgres connection parameters from configs
        self.pg_host = PG_HOST
        self.pg_port = PG_PORT
        self.pg_database = PG_DATABASE
        self.pg_user = PG_USER
        self.pg_password = PG_PASSWORD

        try:
            self.conn = psycopg2.connect(
                host=self.pg_host,
                port=self.pg_port,
                database=self.pg_database,
                user=self.pg_user,
                password=self.pg_password,
            )
            # Register the pgvector adapter so that Python lists are correctly sent as vector type.
            register_vector(self.conn)
            logging.info("Connected to PostgreSQL successfully.")
        except Exception:
            logging.exception("Failed to connect to PostgreSQL.")
            raise

        self.section_headers = {
            "Product Note": [r"Product Note", r"product_note", r"Product\s*Overview"],
            "Objective & Overview": [
                r"Objective\s*&\s*Overview",
                r"Objective\s*and\s*Overview",
            ],
            "Key Functionalities": [r"Key\s*Functionalities"],
            "Pre-Conditions": [r"Pre-Conditions"],
            "Flow": [r"Flow"],
            "Additional Notes": [r"Additional\s*Notes", r"Notes", r"Note"],
        }

    def get_azure_openai_embedding(self, text: str) -> List[float]:
        """Get the embedding for the given text using azure openai text-embedding-3-large (from dim=3072 to dim=768)."""
        return (
            self.azure_openai_embedding_client.embeddings.create(
                input=text, model="text-embedding-3-large", dimensions=768
            )
            .data[0]
            .embedding
        )

    def ensure_collection_exists(self, collection_name: str) -> None:
        """
        Create the table (collection) if it doesn't exist.
        Each table stores:
            - id: UUID primary key
            - vector: pgvector column with dimension 768 (using cosine similarity via the vector_cosine_ops operator)
            - payload: JSONB data containing the document text and metadata.
        An index is created on the vector column.
        """
        create_table_query = f"""
        CREATE TABLE IF NOT EXISTS {collection_name} (
            id UUID PRIMARY KEY,
            vector vector(768),
            payload JSONB
        );
        """
        # Create an index on the vector column using the cosine distance operator.
        # (Adjust the "lists" parameter as needed for your data.)
        create_index_query = f"""
        DO $$
        BEGIN
            IF NOT EXISTS (
                SELECT 1 FROM pg_class c
                JOIN pg_namespace n ON n.oid = c.relnamespace
                WHERE c.relname = '{collection_name}_vector_idx'
            ) THEN
                CREATE INDEX {collection_name}_vector_idx ON {collection_name}
                USING ivfflat (vector vector_cosine_ops) WITH (lists = 100);
            END IF;
        END
        $$;
        """
        try:
            with self.conn.cursor() as cur:
                cur.execute(create_table_query)
                cur.execute(create_index_query)
            self.conn.commit()
            logging.info(f"Ensured collection (table) {collection_name} exists.")
        except Exception:
            self.conn.rollback()
            logging.exception(f"Failed to create collection (table): {collection_name}")

    def delete_old_entries(self, unique_content_id: str, collection_name: str) -> None:
        """
        Remove outdated entries from the table by deleting rows where the payload’s unique_content_id matches.
        """
        delete_query = f"""
        DELETE FROM {collection_name}
        WHERE payload->'metadata'->>'unique_content_id' = %s;
        """
        try:
            with self.conn.cursor() as cur:
                cur.execute(delete_query, (unique_content_id,))
                if cur.rowcount > 0:
                    logging.info(
                        f"Deleted old entries for unique_content_id: {unique_content_id} in {collection_name}."
                    )
            self.conn.commit()
        except Exception:
            self.conn.rollback()
            logging.exception("Failed to delete old entries.")

    def upload_excel_faq_to_vdb(self, excel_file: str, collection_name: str) -> None:
        """
        Upload the FAQ data from the Excel file to the PostgreSQL table.
        Each row is converted to a vector via the embedding endpoint and inserted into the table.
        """
        self.ensure_collection_exists(collection_name)
        filename = os.path.basename(excel_file)
        file_id = str(uuid.uuid4())

        # === READ THE EXCEL FILE ===
        # Read all sheets; keys are sheet names (categories)
        sheets_dict = pd.read_excel(excel_file, sheet_name=None, header=0)

        points = []  # Will hold tuples of (id, vector, payload)
        language = None
        for sheet_name, df in sheets_dict.items():
            df.columns = df.columns.str.strip()

            required_cols_en = {"FAQ Title", "Answer"}
            required_cols_ar = {"Question", "Answer"}
            if not required_cols_en.issubset(
                set(df.columns)
            ) and not required_cols_ar.issubset(set(df.columns)):
                logging.warning(
                    f"Sheet '{sheet_name}' is missing one or more required columns. Found columns: {df.columns.tolist()}"
                )
                continue

            for _, row in df.iterrows():
                if "FAQ Title" in row:
                    title = str(row["FAQ Title"])
                    answer = str(row["Answer"])
                    language = "en"
                else:
                    title = str(row["Question"])
                    answer = str(row["Answer"])
                    language = "ar"

                text_n_ans = f"{title}\n{answer}"
                point_id = str(uuid.uuid4())
                unique_content_id = hashlib.md5(f"{title}".encode()).hexdigest()

                self.delete_old_entries(unique_content_id, collection_name)

                # === Generate the vector embedding for the FAQ Title ===
                vector = self.get_azure_openai_embedding(title)

                # === Prepare metadata ===
                doc = {
                    "text": text_n_ans,
                    "metadata": {
                        "faq_title": title,
                        "answer": answer,
                        "doc_identifier": file_id,
                        "file_name": filename,
                        "category": sheet_name,
                        "unique_content_id": unique_content_id,
                        "language": language,
                    },
                }

                points.append((point_id, vector, json.dumps(doc)))

        # === UPLOAD POINTS TO POSTGRES IN BATCHES ===
        insert_query = f"""
        INSERT INTO {collection_name} (id, vector, payload)
        VALUES %s;
        """
        try:
            with self.conn.cursor() as cur:
                for i in range(0, len(points), QDRANT_UPLOAD_BATCH_SIZE):
                    batch_points = points[i : i + QDRANT_UPLOAD_BATCH_SIZE]
                    psycopg2.extras.execute_values(
                        cur, insert_query, batch_points, template="(%s, %s, %s)"
                    )
                    self.conn.commit()
                    logging.info(
                        f"Uploaded FAQ entries {i + 1} to {i + len(batch_points)}"
                    )
            logging.info("All FAQs successfully inserted into PostgreSQL!")
        except Exception:
            self.conn.rollback()
            logging.exception("Failed to insert FAQ data into PostgreSQL.")

        self.upload_excel_fee_comm_faq_to_vd(
            excel_file, collection_name, file_id, filename, language
        )

    def upload_excel_fee_comm_faq_to_vd(
        self,
        excel_file: str,
        collection_name: str,
        file_id: str,
        filename: str,
        language: str,
    ) -> None:
        """
        Upload the FAQ Fee and Commission data from the Excel file to the PostgreSQL table.
        Each row is converted to a vector via the embedding endpoint and inserted into the table.
        """
        self.ensure_collection_exists(collection_name)

        # === READ THE EXCEL FILE ===
        # Read all sheets; keys are sheet names (categories)
        sheets_dict = pd.read_excel(excel_file, sheet_name=None, header=0)

        points = []  # Will hold tuples of (id, vector, payload)
        for sheet_name, df in sheets_dict.items():
            df.columns = df.columns.str.strip()

            required_cols_en = {"Item", "Type", "Fees/Commissions JOD"}
            required_cols_ar = {"البند", "نوع الحركة", "Fees/Commissions دينار أردني"}
            if not required_cols_en.issubset(
                set(df.columns)
            ) and not required_cols_ar.issubset(set(df.columns)):
                logging.warning(
                    f"Sheet '{sheet_name}' is missing one or more required columns. Found columns: {df.columns.tolist()}"
                )
                continue

            for _, row in df.iterrows():
                if "Item" in row:
                    item = str(row["Item"])
                    type = str(row["Type"])
                    fee = str(row["Fees/Commissions JOD"])
                else:
                    item = str(row["البند"])
                    type = str(row["نوع الحركة"])
                    fee = str(row["Fees/Commissions دينار أردني"])

                if language == "en" or language is None:
                    ques = self.generate_faq_title(item, type)
                    text_n_ans = f"{ques}\n{fee}"
                else:
                    ques = str(row["البند"])
                    text_n_ans = f"{ques} ({type})\n{fee}"

                point_id = str(uuid.uuid4())
                unique_content_id = hashlib.md5(f"{ques}".encode()).hexdigest()

                self.delete_old_entries(unique_content_id, collection_name)

                # === Generate the vector embedding for the FAQ Title ===
                vector = self.get_azure_openai_embedding(ques)

                # === Prepare metadata ===
                doc = {
                    "text": text_n_ans,
                    "metadata": {
                        "faq_title": ques,
                        "answer": fee,
                        "doc_identifier": file_id,
                        "file_name": filename,
                        "category": sheet_name,
                        "unique_content_id": unique_content_id,
                        "language": language,
                    },
                }

                points.append((point_id, vector, json.dumps(doc)))

            formatted_points = self.format_fee_data_to_prompt(
                excel_file,
                sheet_name,
                header="Fee and commissions of Reflect's services",
                language=language,
            )
            vector_aggred = self.get_azure_openai_embedding("Fee and commissions")
            doc_aggred = {
                "text": formatted_points,
                "metadata": {
                    "faq_title": "Fee and commissions of Reflect's services",
                    "answer": formatted_points,
                    "doc_identifier": file_id,
                    "file_name": filename,
                    "category": sheet_name,
                    "unique_content_id": hashlib.md5(
                        "Fee and commissions".encode()
                    ).hexdigest(),
                },
            }
            points.append((str(uuid.uuid4()), vector_aggred, json.dumps(doc_aggred)))

        # === UPLOAD POINTS TO POSTGRES IN BATCHES ===
        insert_query = f"""
        INSERT INTO {collection_name} (id, vector, payload)
        VALUES %s;
        """
        try:
            with self.conn.cursor() as cur:
                for i in range(0, len(points), QDRANT_UPLOAD_BATCH_SIZE):
                    batch_points = points[i : i + QDRANT_UPLOAD_BATCH_SIZE]
                    psycopg2.extras.execute_values(
                        cur, insert_query, batch_points, template="(%s, %s, %s)"
                    )
                    self.conn.commit()
                    logging.info(
                        f"Uploaded FAQ entries {i + 1} to {i + len(batch_points)}"
                    )
            logging.info("All FAQs successfully inserted into PostgreSQL!")
        except Exception:
            self.conn.rollback()
            logging.exception("Failed to insert FAQ data into PostgreSQL.")

    def format_fee_data_to_prompt(
        self,
        file_path,
        sheet_name="Commissions and Fees",
        header="Fee and commissions of Reflect's services",
        language="en",
    ):
        df = pd.read_excel(file_path, sheet_name=sheet_name)

        # Fill NaNs with empty string (optional, for cleanliness)
        df = df.fillna("")

        # Format each row into a single string
        if language == "en":
            rows = [
                f"Item: {row['Item']}, Type: {row['Type']}, Fees/Commissions JOD: {row['Fees/Commissions JOD']}"
                for _, row in df.iterrows()
            ]
        else:
            rows = [
                f"البند: {row['البند']}, نوع الحركة: {row['نوع الحركة']}, Fees/Commissions دينار أردني: {row['Fees/Commissions دينار أردني']}"
                for _, row in df.iterrows()
            ]

        # Combine everything into one big string
        final_str = header + "\n" + "\n".join(rows)
        return final_str

    def generate_faq_title(self, item, type_):
        prompt = f"Generate a clear FAQ question based on the following information:\n\nItem: {item}\nType: {type_}\n\nThe question should ask about fees or commissions. Example: 'What are the fees and commissions that Reflect charges for money transfers using CliQ?', 'What are the fees and commissions that Reflect charges for opening a multicurrency sub-account?', 'What are the fees and commissions that Reflect charges if the minimum balance is not maintained in a Current Account?'"

        try:
            response = self.azure_openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a helpful assistant that generates FAQ questions.",
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.5,
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(
                f"Error generating question for Item '{item}' and Type '{type_}': {e}"
            )
            return f"What are the fees for {type_} {item}?"

    def extract_excel_conversation_history_gist(
        self,
        *,
        df: pd.DataFrame = None,
        excel_file: str = None,
        print_json: bool = False,
        output_path: str = None,
    ) -> List[Dict[str, str]]:
        """
        Reads the Excel file or DataFrame, processes the conversation history, and extracts Q-A pairs.
        Optionally prints the extracted Q-A pairs as JSON.

        Args:
            df: Input DataFrame containing conversation history.
            excel_file: Path to the Excel file containing conversation history.
            print_json: If True, prints the extracted Q-A pairs as JSON.
            output_path: If print_json = True, path to save the extracted Q-A pairs as JSON.

        Returns:
            List of dictionaries containing Q-A pairs.
        """
        # === READ THE EXCEL FILE ===
        if excel_file is not None:
            df = pd.read_excel(excel_file, header=0)
        elif df is None:
            logging.error("Please provide either an Excel file or a DataFrame.")
            return []

        df.columns = df.columns.str.strip().str.lower().str.replace(" ", "_")

        if "createdat" not in df.columns:
            logging.warning(
                f"Skipping file {excel_file}: 'createdat' column not found. Available columns: {df.columns.tolist()}"
            )
            return

        df["createdat"] = pd.to_datetime(df["createdat"], errors="coerce")

        # Filter messages to include only CUSTOMER and AGENT roles
        df = df[df["authortype"].isin(["CUSTOMER", "AGENT"])]

        qa_pairs_list = []  # List to store all Q-A pairs

        # Group messages by conversationid
        for conv_id, group in df.groupby("conversationid"):
            group = group.sort_values("createdat")

            # Build the conversation transcript (each line shows role and message)
            conversation_lines = []
            for _, row in group.iterrows():
                line = f"{row['authortype']}: {row['content'].strip()}"
                conversation_lines.append(line)
            conversation_text = "\n".join(conversation_lines)

            # Use the LLM to extract Q-A pairs
            qa_pairs = self.extract_qa_pairs(conversation_text)
            if qa_pairs is None:
                continue

            for qa in qa_pairs:
                question = qa.get("question", "").strip()
                question_gist = qa.get("question_gist", "").strip() or question
                answer = qa.get("answer", "").strip()
                if not question or not answer:
                    continue

                qa_pairs_list.append(
                    {
                        "conversation_id": conv_id,
                        "question": question,
                        "question_gist": question_gist,
                        "answer": answer,
                        "raw_conversation": conversation_text,
                    }
                )

        # Optionally output the Q-A pairs as JSON
        if print_json:
            with open(output_path, "w", encoding="utf-8") as json_file:
                json.dump(qa_pairs_list, json_file, indent=4, ensure_ascii=False)

        return qa_pairs_list

    def upload_excel_conversation_history_gist_to_vdb(
        self, qa_pairs: Union[List[Dict[str, str]], str], collection_name: str
    ) -> None:
        """
        Uploads the extracted Q-A pairs to PostgreSQL after generating vector embeddings.

        Args:
            qa_pairs: List of dictionaries containing Q-A pairs or path to a JSON file.
            collection_name: Name of the PostgreSQL collection to upload to.
        """
        self.ensure_collection_exists(collection_name)

        # Load from JSON file if qa_pairs is a file path
        if isinstance(qa_pairs, str):
            try:
                with open(qa_pairs, "r", encoding="utf-8") as file:
                    qa_pairs_list = json.load(file)
            except Exception as e:
                logging.exception(f"Error loading JSON file: {e}")
                return
        else:
            qa_pairs_list = qa_pairs

        knowledge_points = []  # List of tuples (id, vector, payload)

        for qa in qa_pairs_list:
            question = qa.get("question", "").strip()
            question_gist = qa.get("question_gist", "").strip() or question
            answer = qa.get("answer", "").strip()
            conv_id = qa.get("conversation_id", "")
            conversation_text = qa.get("raw_conversation", "")

            if not question or not answer:
                continue

            ques_n_ans = f"{question_gist}\n{answer}"
            point_id = str(uuid.uuid4())
            unique_content_id = hashlib.md5(f"{question}".encode()).hexdigest()

            self.delete_old_entries(unique_content_id, collection_name)

            # === Generate the vector embedding for the question ===
            vector = self.get_azure_openai_embedding(question)

            # === Prepare metadata ===
            doc = {
                "text": ques_n_ans,
                "metadata": {
                    "question": question,
                    "answer": answer,
                    "unique_content_id": unique_content_id,
                    "conversation_id": conv_id,
                    "raw_conversation": conversation_text,
                },
            }

            knowledge_points.append((point_id, vector, json.dumps(doc)))

        # === UPLOAD POINTS TO POSTGRES IN BATCHES ===
        insert_query = f"""
        INSERT INTO {collection_name} (id, vector, payload)
        VALUES %s;
        """
        try:
            with self.conn.cursor() as cur:
                for i in range(0, len(knowledge_points), QDRANT_UPLOAD_BATCH_SIZE):
                    batch_points = knowledge_points[i : i + QDRANT_UPLOAD_BATCH_SIZE]
                    psycopg2.extras.execute_values(
                        cur, insert_query, batch_points, template="(%s, %s, %s)"
                    )
                    self.conn.commit()
                    logging.info(
                        f"Uploaded Q-A pairs {i + 1} to {i + len(batch_points)}"
                    )
            logging.info(
                "All conversation-based Q-A pairs successfully inserted into PostgreSQL!"
            )
        except Exception:
            self.conn.rollback()
            logging.exception("Failed to insert conversation Q-A data into PostgreSQL.")

    def extract_and_upload_excel_conversation_history_gist_to_vdb(
        self, *, df: pd.DataFrame = None, excel_file: str = None, collection_name: str
    ) -> None:
        """
        Upload the conversation history gist data from the Excel file or DataFrame to PostgreSQL.
        Each conversation is grouped, processed via the LLM to extract Q-A pairs, vectorized, and then inserted.
        """
        self.ensure_collection_exists(collection_name)

        # === READ THE EXCEL FILE ===
        if excel_file is not None:
            df = pd.read_excel(excel_file, header=0)
        elif df is None:
            logging.error("Please provide either an Excel file or a DataFrame.")
            return

        df.columns = df.columns.str.strip().str.lower().str.replace(" ", "_")

        if "createdat" not in df.columns:
            logging.info(
                f"Skipping file {excel_file}: 'createdat' column not found. Available columns: {df.columns.tolist()}"
            )
            return

        df["createdat"] = pd.to_datetime(df["createdat"], errors="coerce")

        # Filter messages to include only CUSTOMER and AGENT roles
        df = df[df["authortype"].isin(["CUSTOMER", "AGENT"])]

        knowledge_points = []  # List of tuples (id, vector, payload)

        # Group messages by conversationid
        for conv_id, group in df.groupby("conversationid"):
            group = group.sort_values("createdat")

            # Build the conversation transcript (each line shows role and message)
            conversation_lines = []
            for _, row in group.iterrows():
                line = f"{row['authortype']}: {row['content'].strip()}"
                conversation_lines.append(line)
            conversation_text = "\n".join(conversation_lines)

            # Use the LLM to extract Q-A pairs
            qa_pairs = self.extract_qa_pairs(conversation_text)
            if qa_pairs is None:
                continue

            for qa in qa_pairs:
                question = qa.get("question", "").strip()
                question_gist = qa.get("question_gist", "").strip() or question
                answer = qa.get("answer", "").strip()
                if not question or not answer:
                    continue

                ques_n_ans = f"{question_gist}\n{answer}"
                point_id = str(uuid.uuid4())
                unique_content_id = hashlib.md5(f"{question}".encode()).hexdigest()

                self.delete_old_entries(unique_content_id, collection_name)

                # === Generate the vector embedding for the question ===
                vector = self.get_azure_openai_embedding(question)

                # === Prepare metadata ===
                doc = {
                    "text": ques_n_ans,
                    "metadata": {
                        "question": question,
                        "answer": answer,
                        "unique_content_id": unique_content_id,
                        "conversation_id": conv_id,
                        "raw_conversation": conversation_text,
                    },
                }

                knowledge_points.append((point_id, vector, json.dumps(doc)))

        # === UPLOAD POINTS TO POSTGRES IN BATCHES ===
        insert_query = f"""
        INSERT INTO {collection_name} (id, vector, payload)
        VALUES %s;
        """
        try:
            with self.conn.cursor() as cur:
                for i in range(0, len(knowledge_points), QDRANT_UPLOAD_BATCH_SIZE):
                    batch_points = knowledge_points[i : i + QDRANT_UPLOAD_BATCH_SIZE]
                    psycopg2.extras.execute_values(
                        cur, insert_query, batch_points, template="(%s, %s, %s)"
                    )
                    self.conn.commit()
                    logging.info(
                        f"Uploaded Q-A pairs {i + 1} to {i + len(batch_points)}"
                    )
            logging.info(
                "All conversation-based Q-A pairs successfully inserted into PostgreSQL!"
            )
        except Exception:
            self.conn.rollback()
            logging.exception("Failed to insert conversation Q-A data into PostgreSQL.")

    def extract_qa_pairs(
        self, conversation_text: str, print_json: bool = False
    ) -> List[Dict[str, str]]:
        """
        Extract Q-A pairs from a conversation transcript using LLM.

        Instructions:
        - If the conversation discusses one main FAQ, merge all relevant CUSTOMER messages into a single question and
            all corresponding AGENT replies into a single answer.
        - If the conversation contains multiple distinct Q-A pairs, extract each pair separately.
        - If no meaningful Q-A pair exists (e.g., just greetings or small talk), output "null".

        The LLM should return valid JSON. If multiple pairs exist, it should return an array of objects with keys
        "question" and "answer". If there is only one pair, it may return either a single JSON object or an array
        containing one object.
        """
        prompt = f"""
        You are an assistant that analyzes a conversation transcript between a CUSTOMER and an AGENT.
        Extract Q-A pairs and ensure JSON output. 
        The transcript contains multiple messages, each labeled with its sender ("CUSTOMER" or "AGENT").
        Sometimes several messages from the CUSTOMER and several replies from the AGENT address the same FAQ;
        in that case, merge them into one aggregated Q-A pair.
        However, if the conversation contains multiple distinct Q-A pairs (discussing different topics), extract each pair separately without merging them.
        For each Q-A pair, also provide a concise summary (a "gist") of the question using key phrases.
        For example, if the aggregated question is "How do I reset my password? I forgot my current password too.", the gist might be "reset password, forgot password".
        If the conversation does not contain any meaningful Q-A pair (for example, if it only contains greetings or small talk), output "null".
        If multiple Q-A pairs exist, return an array of objects, each with keys "question", "question_gist", and "answer".
        If there is only one pair, you may return either a single JSON object or an array with one object.

        Return your response strictly as valid JSON, enclosed within triple backticks (```json ... ```).

        Conversation transcript:
        {conversation_text}
        """
        try:
            response = self.azure_openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                max_tokens=1000,
            )

            if (
                response is None
                or not hasattr(response, "choices")
                or not response.choices
            ):
                logging.exception("OpenAI API did not return a valid response.")
                return None

            answer_text = (
                response.choices[0].message.content.strip()
                if response.choices[0].message.content
                else ""
            )

            if not answer_text:
                logging.exception("LLM returned an empty response.")
                return None

            # Extract JSON content if wrapped in triple backticks
            import re

            json_match = re.search(r"```json\n(.*?)\n```", answer_text, re.DOTALL)
            if json_match:
                answer_text = json_match.group(1)

            try:
                parsed = json.loads(answer_text)
            except json.JSONDecodeError as e:
                logging.exception(
                    f"Invalid JSON output from LLM: {answer_text}, Error: {str(e)}"
                )
                return None

            # Ensure output is a list
            qa_pairs = (
                [parsed]
                if isinstance(parsed, dict)
                else parsed
                if isinstance(parsed, list)
                else None
            )

            if print_json and qa_pairs:
                with open("qa_pairs.json", "w") as json_file:
                    json.dump(qa_pairs, json_file, indent=4)

            return qa_pairs
        except Exception as e:
            logging.exception(f"Error processing conversation via LLM: {str(e)}")
            return None

    def upload_product_notes_to_vdb(
        self, *, pdf: str = None, collection_name: str = None
    ) -> None:
        """
        Upload the product notes from the PDF file to the PostgreSQL table.
        The PDF is partitioned into sections and each section is vectorized and inserted.
        """
        self.ensure_collection_exists(collection_name)
        chunks = self.extract_product_notes(pdf)

        doc_identifier = str(uuid.uuid4())

        points = []
        for _, chunk in enumerate(chunks):
            product_note = chunk["product_note"]
            section = chunk["section"]
            content = chunk["content"]
            full_text = chunk["full_text"]

            point_id = str(uuid.uuid4())
            unique_content_id = hashlib.md5(f"{full_text}".encode()).hexdigest()

            self.delete_old_entries(unique_content_id, collection_name)

            # === Generate the vector embedding for the product note ===
            vector = self.get_azure_openai_embedding(full_text)

            # === Prepare metadata ===
            doc = {
                "text": full_text,
                "metadata": {
                    "product_note": product_note,
                    "section": section,
                    "content": content,
                    "doc_identifier": doc_identifier,
                    "unique_content_id": unique_content_id,
                },
            }

            points.append((point_id, vector, json.dumps(doc)))

        # === UPLOAD POINTS TO POSTGRES IN BATCHES ===
        insert_query = f"""
        INSERT INTO {collection_name} (id, vector, payload)
        VALUES %s;
        """
        try:
            with self.conn.cursor() as cur:
                for i in range(0, len(points), QDRANT_UPLOAD_BATCH_SIZE):
                    batch_points = points[i : i + QDRANT_UPLOAD_BATCH_SIZE]
                    psycopg2.extras.execute_values(
                        cur, insert_query, batch_points, template="(%s, %s, %s)"
                    )
                    self.conn.commit()
                    logging.info(
                        f"Uploaded product notes {i + 1} to {i + len(batch_points)}"
                    )
            logging.info("All product notes successfully inserted into PostgreSQL!")
        except Exception:
            self.conn.rollback()
            logging.exception("Failed to insert product notes into PostgreSQL.")

        df = self.process_single_pdf(pdf)
        self.upload_faq_df_to_vdb(df, pdf, collection_name, doc_identifier)

    # ------------NEW FUNCTIONS START------------
    def upload_points_to_postgres(self, points, collection_name: str):
        insert_query = f"""
        INSERT INTO {collection_name} (id, vector, payload)
        VALUES %s;
        """
        try:
            with self.conn.cursor() as cur:
                for i in range(0, len(points), QDRANT_UPLOAD_BATCH_SIZE):
                    batch_points = points[i : i + QDRANT_UPLOAD_BATCH_SIZE]
                    psycopg2.extras.execute_values(
                        cur, insert_query, batch_points, template="(%s, %s, %s)"
                    )
                    self.conn.commit()
                    logging.info(
                        f"Uploaded FAQ entries {i + 1} to {i + len(batch_points)}"
                    )
            logging.info("All FAQs successfully inserted into PostgreSQL!")
        except Exception:
            self.conn.rollback()
            logging.exception("Failed to insert FAQ data into PostgreSQL.")

    def remove_arabic(self, docx_path: str) -> str:
        """
        Extracts English portion from DOCX, then processes into chunks.
        Assumes English content appears first, followed by Arabic.
        """
        document = Document(docx_path)
        full_text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        # Heuristic: Cut the text before Arabic starts
        # Arabic Unicode range is \u0600-\u06FF
        arabic_start = re.search(r"[\u0600-\u06FF]", full_text)
        if arabic_start:
            full_text = full_text[: arabic_start.start()]
        return full_text

    def split_english_arabic(self, docx_path: str) -> tuple[str, str]:
        """
        Separates English and Arabic portions from a DOCX file.
        Assumes English appears first, followed by Arabic.
        """
        document = Document(docx_path)
        full_text = "\n".join(p.text for p in document.paragraphs if p.text.strip())

        # Arabic Unicode range is \u0600-\u06FF
        arabic_start = re.search(r"[\u0600-\u06FF]", full_text)

        if arabic_start:
            english_text = full_text[: arabic_start.start()].strip()
            arabic_text = full_text[arabic_start.start() :].strip()
        else:
            english_text = full_text.strip()
            arabic_text = ""

        return english_text, arabic_text

    def get_filename(self, path: str) -> str:
        filename = os.path.basename(path)
        return filename

    def upload_product_notes_v2_to_vdb(
        self, *, docx: str = None, collection_name: str = None
    ) -> None:
        """
        Upload the product notes from the PDF file to the PostgreSQL table.
        Preprocess Word to remove arabic text and convert it to PDF.
        The PDF is partitioned into sections and each section is vectorized and inserted.

        Filename expected to be something like '12313313New Filename.docx'
        """
        self.ensure_collection_exists(collection_name)

        file_name = self.get_filename(docx)

        processed_file_name = file_name.split(" ", 1)

        texts = self.remove_arabic(docx)

        faqs = self.generate_faq_pairs(texts)

        df = pd.DataFrame()

        if faqs:
            df = pd.DataFrame(faqs)
        else:
            print(f"No FAQs generated for {file_name}")

        file_name_vector = self.get_azure_openai_embedding(processed_file_name)

        full_content_vector = self.get_azure_openai_embedding(texts)

        points = []

        unique_content_id = hashlib.md5(f"{texts}".encode()).hexdigest()

        doc_identifier = str(uuid.uuid4())

        doc = {
            "text": texts,
            "metadata": {
                "file_name": file_name,
                "content": texts,
                "doc_identifier": doc_identifier,
                "unique_content_id": unique_content_id,
            },
        }

        points.append((str(uuid.uuid4()), file_name_vector, json.dumps(doc)))

        points.append((str(uuid.uuid4()), full_content_vector, json.dumps(doc)))

        self.upload_points_to_postgres(points, collection_name)

        self.upload_faq_df_to_vdb(df, docx, collection_name, doc_identifier)

    # ------------NEW FUNCTIONS END------------

    def upload_product_notes_to_vdb_v3(
        self, *, docx: str = None, collection_name: str = None
    ) -> None:
        """
        Upload the product notes from the PDF file to the PostgreSQL table.
        Preprocess Word to remove arabic text and convert it to PDF.
        The PDF is partitioned into sections and each section is vectorized and inserted.

        Filename expected to be something like '12313313New Filename.docx'
        """
        self.ensure_collection_exists(collection_name)

        file_name = self.get_filename(docx)

        processed_file_name = file_name.split(" ", 1)

        en_text, ar_text = self.split_english_arabic(docx)
        combined_text = f"{en_text}\n\n{ar_text}".strip()

        en_faqs = self.generate_faq_pairs(en_text)
        ar_faqs = self.generate_faq_pairs(ar_text)

        en_df = pd.DataFrame()
        ar_df = pd.DataFrame()

        if en_faqs:
            en_df = pd.DataFrame(en_faqs)
        if ar_faqs:
            ar_df = pd.DataFrame(ar_faqs)
        else:
            print(f"No FAQs generated for {file_name}")

        file_name_vector = self.get_azure_openai_embedding(processed_file_name)

        en_full_content_vector = self.get_azure_openai_embedding(en_text)
        ar_full_content_vector = self.get_azure_openai_embedding(ar_text)

        points = []

        en_unique_content_id = hashlib.md5(f"{en_text}".encode()).hexdigest()
        ar_unique_content_id = hashlib.md5(f"{ar_text}".encode()).hexdigest()
        mix_unique_content_id = hashlib.md5(f"{combined_text}".encode()).hexdigest()

        doc_identifier = str(uuid.uuid4())

        full_doc = {
            "text": combined_text,
            "metadata": {
                "file_name": file_name,
                "content": combined_text,
                "language": "mix",
                "doc_identifier": doc_identifier,
                "unique_content_id": mix_unique_content_id,
            },
        }

        en_doc = {
            "text": en_text,
            "metadata": {
                "file_name": file_name,
                "content": en_text,
                "language": "en",
                "doc_identifier": doc_identifier,
                "unique_content_id": en_unique_content_id,
            },
        }

        ar_doc = {
            "text": ar_text,
            "metadata": {
                "file_name": file_name,
                "content": ar_text,
                "language": "ar",
                "doc_identifier": doc_identifier,
                "unique_content_id": ar_unique_content_id,
            },
        }

        points.append((str(uuid.uuid4()), file_name_vector, json.dumps(full_doc)))

        points.append((str(uuid.uuid4()), en_full_content_vector, json.dumps(en_doc)))

        points.append((str(uuid.uuid4()), ar_full_content_vector, json.dumps(ar_doc)))

        self.upload_points_to_postgres(points, collection_name)

        self.upload_faq_df_to_vdb(en_df, docx, collection_name, doc_identifier)

        self.upload_faq_df_to_vdb(ar_df, docx, collection_name, doc_identifier)

    def extract_product_notes(self, pdf: str) -> List[Dict]:
        """
        Partition PDF and extract sections, handling both standard and outlier documents.
        """
        elements = partition_pdf(
            pdf,
            strategy="auto",
            include_element_types=[ElementType.TEXT, ElementType.TABLE],
        )

        # Extract full text for regex processing.
        full_text = "\n".join([el.text for el in elements if el.text])

        # Try standard processing first.
        standard_chunks = self._process_standard_pdf(full_text, pdf)

        # Count how many chunks match our predefined standard headers.
        num_matched = sum(
            1 for chunk in standard_chunks if chunk["section"] in self.section_headers
        )
        if num_matched > 3:
            return standard_chunks

        # If fewer than four sections match, re-partition using by_title chunking strategy.
        elements_by_title = partition_pdf(
            pdf,
            chunking_strategy="by_title",
            new_after_n_chars=300,
            include_element_types=[ElementType.TEXT, ElementType.TABLE],
        )
        return self._process_outlier_pdf(elements_by_title, pdf)

    def _process_standard_pdf(self, full_text: str, pdf: str) -> List[Dict]:
        """Process PDF using standard section headers."""
        product_note_pattern = (
            r"^Product Note:\s*(.*?)\s*(?=("
            + "|".join(
                [
                    variant
                    for variations in self.section_headers.values()
                    for variant in variations
                ]
            )
            + "))"
        )

        product_note_match = re.search(
            product_note_pattern, full_text, re.MULTILINE | re.IGNORECASE
        )

        if product_note_match:
            product_note = f"Product Note: {product_note_match.group(1).strip()}"
        else:
            pdf_name = os.path.splitext(os.path.basename(pdf))[0]
            product_note = f"Product Note: {pdf_name}"

        parts = re.split(
            r"("
            + "|".join(
                [
                    variant
                    for variations in self.section_headers.values()
                    for variant in variations
                ]
            )
            + ")",
            full_text,
            flags=re.IGNORECASE,
        )

        chunks = []
        for i in range(1, len(parts), 2):
            header = parts[i].strip()
            content = parts[i + 1].strip() if i + 1 < len(parts) else ""

            # Standardize the header.
            for key, variations in self.section_headers.items():
                if any(
                    re.fullmatch(variant, header, re.IGNORECASE)
                    for variant in variations
                ):
                    header = key
                    break

            # Remove (cid:127) from header and content.
            header = re.sub(r"\(cid:127\)", "", header, flags=re.IGNORECASE)
            cleaned_content = re.sub(r"\(cid:127\)", "", content, flags=re.IGNORECASE)

            if cleaned_content:
                chunk_text = f"{product_note} - {header}\n{cleaned_content}"
                chunks.append(
                    {
                        "product_note": product_note,
                        "section": header,
                        "content": cleaned_content,
                        "full_text": chunk_text,
                    }
                )

        return chunks

    def _process_outlier_pdf(self, elements: List, pdf: str) -> List[Dict]:
        """Process PDFs with non-standard layouts using the by_title chunking strategy."""
        pdf_name = os.path.splitext(os.path.basename(pdf))[0]

        # Filter relevant elements: those with nonempty text and a type in ["Text", "Table"].
        relevant_elements = [
            el
            for el in elements
            if hasattr(el, "text")
            and el.text.strip()
            and getattr(el, "type", "Text") in ["Text", "Table"]
        ]

        chunks = []
        for el in relevant_elements:
            # Split text into nonempty lines.
            lines = [line.strip() for line in el.text.splitlines() if line.strip()]
            if not lines:
                continue

            header = lines[0]

            # Remove (cid:127) from header.
            header = re.sub(r"\(cid:127\)", "", header, flags=re.IGNORECASE)

            # Check if the header is duplicated as the first non-header line.
            content_lines = lines[1:]
            if content_lines and content_lines[0] == header:
                content_lines = content_lines[1:]
            content = "\n".join(content_lines)

            # Remove (cid:127) from content.
            content = re.sub(r"\(cid:127\)", "", content, flags=re.IGNORECASE)

            # Compose full_text; if no additional content, use only the header.
            if content:
                full_text = f"{pdf_name} - {header}\n{content}"
            else:
                full_text = f"{pdf_name} - {header}"

            chunks.append(
                {
                    "product_note": pdf_name,
                    "section": header,
                    "content": content,
                    "full_text": full_text,
                }
            )
        return chunks

    def upload_faq_df_to_vdb(
        self,
        faq_df: pd.DataFrame,
        pdf_path: str,
        collection_name: str,
        doc_identifier: str,
    ) -> None:
        """
        Upload the FAQ data from a DataFrame to the PostgreSQL table.
        Each row is converted to a vector via the embedding endpoint and inserted into the table.
        """
        self.ensure_collection_exists(collection_name)

        # Extract filename without extension from the PDF path
        category_name = os.path.splitext(os.path.basename(pdf_path))[0]

        points = []  # Will hold tuples of (id, vector, payload)

        # Ensure required columns exist in the DataFrame
        required_cols = {"FAQ Title", "Answer"}
        if not required_cols.issubset(set(faq_df.columns)):
            logging.warning(
                f"DataFrame is missing one or more required columns. Found columns: {faq_df.columns.tolist()}"
            )
            return

        for _, row in faq_df.iterrows():
            title = str(row["FAQ Title"])
            answer = str(row["Answer"])

            text_n_ans = f"{title}\n{answer}"
            point_id = str(uuid.uuid4())
            unique_content_id = hashlib.md5(f"{title}".encode()).hexdigest()

            self.delete_old_entries(unique_content_id, collection_name)

            # Generate the vector embedding for the FAQ Title
            vector = self.get_azure_openai_embedding(title)

            # Prepare metadata
            doc = {
                "text": text_n_ans,
                "metadata": {
                    "faq_title": title,
                    "doc_identifier": doc_identifier,
                    "answer": answer,
                    "category": category_name,  # Replaces sheet name with filename
                    "unique_content_id": unique_content_id,
                },
            }

            points.append((point_id, vector, json.dumps(doc)))

        # === UPLOAD POINTS TO POSTGRES IN BATCHES ===
        insert_query = f"""
        INSERT INTO {collection_name} (id, vector, payload)
        VALUES %s;
        """
        try:
            with self.conn.cursor() as cur:
                for i in range(0, len(points), QDRANT_UPLOAD_BATCH_SIZE):
                    batch_points = points[i : i + QDRANT_UPLOAD_BATCH_SIZE]
                    psycopg2.extras.execute_values(
                        cur, insert_query, batch_points, template="(%s, %s, %s)"
                    )
                    self.conn.commit()
                    logging.info(
                        f"Uploaded FAQ entries {i + 1} to {i + len(batch_points)}"
                    )
            logging.info("All FAQs successfully inserted into PostgreSQL!")
        except Exception:
            self.conn.rollback()
            logging.exception("Failed to insert FAQ data into PostgreSQL.")

    def extract_text_from_pdf(self, pdf_path):
        """Extracts text from a PDF file."""
        text = ""
        with fitz.open(pdf_path) as doc:
            for page in doc:
                text += page.get_text("text") + "\n"
        return text

    def generate_faq_pairs(self, text):
        """Generates 10 question-answer pairs using Azure OpenAI LLM."""
        prompt = (
            "Generate 10 frequently asked questions (FAQ) and their answers from the following text:\n\n"
            f"{text}\n\n"
            "Format the output as JSON: [{'FAQ Title': 'Question1', 'Answer': 'Answer1'}, ...]"
        )

        try:
            response = self.azure_openai_client.chat.completions.create(
                model=self.azure_openai_deployment_name,
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.7,
            )

            if (
                not response
                or not response.choices
                or not response.choices[0].message.content
            ):
                print("Error: Empty response from OpenAI.")
                return []

            # Extract JSON from response (remove markdown formatting)
            raw_content = response.choices[0].message.content.strip()
            cleaned_content = re.sub(r"```json|```", "", raw_content).strip()

            # Parse the cleaned JSON
            faq_data = json.loads(cleaned_content)
            return faq_data if isinstance(faq_data, list) else []

        except json.JSONDecodeError as e:
            print(f"Error parsing response: {e}")
            return []
        except Exception as e:
            print(f"API Error: {e}")
            return []

    def process_single_pdf(self, pdf_path):
        """Processes a single PDF file, extracts text, generates FAQ pairs, and returns a DataFrame."""
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"The file {pdf_path} does not exist.")

        pdf_name = os.path.splitext(os.path.basename(pdf_path))[
            0
        ]  # Extract filename without extension
        print(f"Processing {pdf_name}...")

        text = self.extract_text_from_pdf(pdf_path)
        faqs = self.generate_faq_pairs(text)

        if faqs:
            df = pd.DataFrame(faqs)
            return df
        else:
            print(f"No FAQs generated for {pdf_name}")
            return pd.DataFrame()


def parse_gcp_url(gcp_url):
    parsed_url = urlparse(gcp_url)
    decoded_path = unquote(parsed_url.path)
    if parsed_url.netloc == "storage.googleapis.com":
        path_parts = decoded_path.lstrip("/").split("/", 1)
        if len(path_parts) == 0 or not path_parts[0]:
            raise ValueError(f"Invalid GCP URL: {gcp_url}")
        bucket_name = path_parts[0]
        object_path = path_parts[1] if len(path_parts) > 1 else ""
    else:
        bucket_name = parsed_url.netloc.split(".")[0]
        object_path = decoded_path.lstrip("/")

    return bucket_name, object_path


def list_gcp_files(bucket_name, project_name, folder_path=""):
    credentials = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE
    )
    storage_client = storage.Client(credentials=credentials, project=project_name)
    bucket = storage_client.bucket(bucket_name)
    blobs = bucket.list_blobs(prefix=folder_path)
    file_urls = [
        f"gs://{bucket_name}/{blob.name}"
        for blob in blobs
        if not blob.name.endswith("/")
    ]
    return file_urls


def download_from_gcp(gcp_url, project_name):
    bucket_name, object_path = parse_gcp_url(gcp_url)
    credentials = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE
    )
    storage_client = storage.Client(credentials=credentials, project=project_name)
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(object_path)
    temp_dir = tempfile.mkdtemp()
    local_filename = os.path.join(temp_dir, os.path.basename(object_path))
    blob.download_to_filename(local_filename)
    logging.info(f"Downloaded {gcp_url} to {local_filename}")
    return local_filename, temp_dir


def main():
    parser = argparse.ArgumentParser(description="Run knowledge upload functions.")
    parser.add_argument(
        "--mode",
        choices=["faq", "conversation", "product", "product_v2", "product_v3"],
        required=True,
    )
    parser.add_argument(
        "--module", choices=["pre_login", "post_login"], default="pre_login"
    )
    parser.add_argument("--project", help="GCP project name.")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--path", help="Local file or folder path.")
    group.add_argument("--storage", help="GCP Storage URL of the file or folder.")
    parser.add_argument("--collection", help="Table name in Postgres.", required=True)
    args = parser.parse_args()

    knowledge_injector = Knowledge(module=args.module)
    temp_dirs = []
    files_to_process = []

    if args.storage:
        bucket_name, object_path = parse_gcp_url(args.storage)
        gcp_files = list_gcp_files(bucket_name, args.project, object_path)

        for file_url in gcp_files:
            if args.mode == "product" and file_url.endswith(".pdf"):
                local_file, temp_dir = download_from_gcp(file_url, args.project)
                files_to_process.append(local_file)
                temp_dirs.append(temp_dir)
            elif args.mode in ["faq", "conversation"] and file_url.endswith(
                (".xls", ".xlsx")
            ):
                local_file, temp_dir = download_from_gcp(file_url, args.project)
                files_to_process.append(local_file)
                temp_dirs.append(temp_dir)
            # TODO - add gcp for new knowledge
    else:
        if os.path.isdir(args.path):
            if args.mode in ["faq", "conversation"]:
                files_to_process.extend(glob.glob(os.path.join(args.path, "*.xls")))
                files_to_process.extend(glob.glob(os.path.join(args.path, "*.xlsx")))
            elif args.mode == "product":
                files_to_process.extend(glob.glob(os.path.join(args.path, "*.pdf")))
            elif args.mode in ["product_v2"]:
                files_to_process.extend(glob.glob(os.path.join(args.path, "*.docx")))
            elif args.mode in ["product_v3"]:
                files_to_process.extend(glob.glob(os.path.join(args.path, "*.docx")))

        elif os.path.isfile(args.path):
            files_to_process = [args.path]
        else:
            logging.error("Specified path does not exist.")
            sys.exit(1)

    if not files_to_process:
        logging.warning("No files found to process.")
        sys.exit(0)

    for file in files_to_process:
        if args.mode == "faq":
            knowledge_injector.upload_excel_faq_to_vdb(
                excel_file=file, collection_name=args.collection
            )
        elif args.mode == "conversation":
            knowledge_injector.extract_excel_conversation_history_gist(
                excel_file=file, print_json=True, output_path="conversation.json"
            )
        elif args.mode == "product":
            knowledge_injector.upload_product_notes_to_vdb(
                pdf=file, collection_name=args.collection
            )
        elif args.mode == "product_v2":
            knowledge_injector.upload_product_notes_v2_to_vdb(
                docx=file, collection_name=args.collection
            )
        elif args.mode == "product_v3":
            knowledge_injector.upload_product_notes_to_vdb_v3(
                docx=file, collection_name=args.collection
            )

    logging.info("Processing completed.")

    for temp_dir in temp_dirs:
        try:
            for file in os.listdir(temp_dir):
                os.remove(os.path.join(temp_dir, file))
            os.rmdir(temp_dir)
            logging.info(f"Temporary files cleaned up: {temp_dir}")
        except Exception as e:
            logging.error(f"Error cleaning up temporary files: {str(e)}")


if __name__ == "__main__":
    main()
